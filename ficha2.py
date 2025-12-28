# -*- coding: utf-8 -*-
"""
ficha2.py — Generador profesional de informes DOCX
Formato Corporativo Global Talent - Version "Todo Terreno" (Bordes + Detección Inteligente)
"""

import os
import re
import json
import argparse
from pathlib import Path
from datetime import datetime

# OpenAI
try:
    from dotenv import load_dotenv
    load_dotenv()
except:
    pass

_HAS_OPENAI = True
try:
    from openai import OpenAI
except:
    _HAS_OPENAI = False

# DOCX
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Extra extraction libs
try:
    import docx2txt
    DOCTXTXT_AVAILABLE = True
except:
    DOCTXTXT_AVAILABLE = False

try:
    from pdfminer.high_level import extract_text as pdf_extract_text
    PDF_AVAILABLE = True
except:
    PDF_AVAILABLE = False


# ===========================================
#          HELPERS
# ===========================================

def read_text(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="ignore")
    except:
        return ""


def read_cv_any(path: Path) -> str:
    if not path or not path.exists():
        return ""
    suf = path.suffix.lower()

    if suf == ".docx" and DOCTXTXT_AVAILABLE:
        return docx2txt.process(str(path)) or ""

    if suf == ".pdf" and PDF_AVAILABLE:
        return pdf_extract_text(str(path)) or ""

    return read_text(path)

# --- DIBUJAR BORDES (ESTILO LUIS BARRIOS) ---
def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.append(tblPr)
    
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        borders.append(border)
    
    tblPr.append(borders)


# ===========================================
#          HERRAMIENTAS (%) - CORREGIDO
# ===========================================

def extract_tools_with_percent(text: str):
    text = text.replace("\r", "")
    # Busca nombre + numero + %
    regex = re.compile(r"(.+?)\s*[:\-]?\s*(\d{1,3})\s*%", re.I)
    tools = []
    for line in text.splitlines():
        m = regex.search(line)
        if m:
            # Limpieza: quita paréntesis abiertos al final
            name = m.group(1).strip().strip(" (")
            perc = m.group(2).strip() + "%"
            tools.append({"herramienta": name, "nivel": perc})
    return tools


# ===========================================
#          NORMALIZADOR (CORE LOGIC)
# ===========================================

def normalize_skills(data: dict) -> dict:
    """
    Revisa si la IA puso software en 'competencias' y lo mueve a 'herramientas'.
    """
    software_keywords = [
        "excel", "word", "powerpoint", "office", "photoshop", "illustrator", 
        "figma", "canva", "trello", "asana", "jira", "python", "java", 
        "javascript", "react", "angular", "node", "html", "css", "sql", 
        "salesforce", "sap", "hubspot", "vs code", "visual studio", "postman",
        "git", "github", "docker", "aws", "azure"
    ]

    comps = data.get("competencias_tecnicas", [])
    tools = data.get("herramientas", [])
    
    new_comps = []
    
    for item in comps:
        nombre = item.get("competencia", "").lower()
        es_software = any(sw in nombre for sw in software_keywords)
        
        if es_software:
            # MOVER A HERRAMIENTAS
            exists = any(t.get("herramienta", "").lower() == nombre for t in tools)
            if not exists:
                tools.append({
                    "herramienta": item.get("competencia"), 
                    "nivel": item.get("nivel", "Se menciona")
                })
        else:
            # MANTENER EN COMPETENCIAS
            new_comps.append(item)
            
    data["competencias_tecnicas"] = new_comps
    data["herramientas"] = tools
    return data


# ===========================================
#          AI (CEREBRO)
# ===========================================

def ai_enabled():
    return _HAS_OPENAI and bool(os.environ.get("OPENAI_API_KEY"))


def ai_extract_fields(full_text, cv_text, tools, model="gpt-4o-mini"):
    if not ai_enabled():
        raise RuntimeError("OPENAI_API_KEY no configurada.")

    if len(cv_text) > 45000:
        cv_text = cv_text[:45000] + "\n...[TRUNCADO]..."

    client = OpenAI()

    # PROMPT MEJORADO: Instrucciones para CVs sin porcentajes
    system = (
        "Eres analista de RRHH. Genera un informe profesional en ESPAÑOL NEUTRO.\n"
        "\n"
        "Devuelve solo JSON válido. Estructura estricta:\n"
        "{\n"
        '  \"fecha\": \"...\",\n'
        '  \"puesto\": \"...\",\n'
        '  \"resumen_ejecutivo\": \"...\",\n'
        '  \"ficha_tecnica\": {\n'
        '       \"nombre\": \"...\",\n'
        '       \"ubicacion\": \"...\",\n'
        '       \"nivel_experiencia\": \"...\",\n'
        '       \"formacion_formal\": \"...\",\n'
        '       \"nivel_ingles\": \"...\",\n'
        '       \"disponibilidad\": \"...\"\n'
        "  },\n"
        '  \"competencias_tecnicas\": [{\"competencia\": \"...\", \"nivel\": \"...\"}],\n'
        '  \"habilidades_blandas\": [{\"habilidad\": \"...\", \"nivel\": \"...\"}],\n'
        '  \"herramientas\": [{\"herramienta\": \"...\", \"nivel\": \"...\"}],\n'
        '  \"plus\": \"...\",\n'
        '  \"formacion_sugerida\": \"...\",\n'
        '  \"recomendacion_final\": \"...\",\n'
        '  \"responsable\": \"Departamento de Recursos Humanos\"\n'
        "}\n\n"
        "REGLAS CRÍTICAS:\n"
        "1. 'competencias_tecnicas': ÁREAS DE CONOCIMIENTO (Ej: 'Gestión de Proyectos', 'Contabilidad', 'Desarrollo Web').\n"
        "   - PROHIBIDO poner nombres de software aquí.\n"
        "2. 'herramientas': SOLO SOFTWARE y PLATAFORMAS (Ej: 'Excel', 'Python', 'Jira').\n"
        "   - IMPORTANTE: Si la lista de herramientas detectadas está vacía o incompleta, EXTRAE LAS HERRAMIENTAS DEL TEXTO DEL CV.\n"
        "   - Si no tienen porcentaje explícito, usa 'Se menciona' o infiere el nivel (Básico/Intermedio/Avanzado) según el contexto.\n"
        "3. 'habilidades_blandas': Soft skills.\n"
    )

    payload = {
        "texto": full_text,
        "cv": cv_text,
        "herramientas_detectadas": tools
    }

    try:
        response = client.responses.create(
            model=model,
            temperature=0.30,
            input=[
                {"role": "system", "content": system},
                {"role": "user", "content": json.dumps(payload, ensure_ascii=False)}
            ]
        )
        raw = response.output_text.strip()
    except AttributeError:
        response = client.chat.completions.create(
            model=model,
            temperature=0.30,
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": json.dumps(payload, ensure_ascii=False)}
            ]
        )
        raw = response.choices[0].message.content.strip()

    raw = re.sub(r"^```json|```$", "", raw).strip()

    try:
        data = json.loads(raw)
    except json.JSONDecodeError:
        data = {}

    # Si la regex encontró algo, le damos prioridad, sino confiamos en la IA
    if tools:
        # Fusionar herramientas de IA con las de Regex si es necesario, 
        # pero por ahora, si regex halló algo, lo usamos.
        data["herramientas"] = tools + [t for t in data.get("herramientas", []) if t["herramienta"] not in [x["herramienta"] for x in tools]]

    return data


# ===========================================
#          RENDER DOCX
# ===========================================

def set_run_font(run, bold=False, size=11):
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    rpr = run._element.rPr
    if rpr is None: return
    rFonts = OxmlElement("w:rFonts")
    for a in ["ascii", "hAnsi", "eastAsia", "cs"]:
        rFonts.set(qn(f"w:{a}"), "Times New Roman")
    rpr.append(rFonts)


def add_paragraph(doc, text, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  before=0, after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    set_run_font(run, bold=bold, size=size)
    return p


def add_label_value(doc, label, val):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    r1 = p.add_run(f"{label}: ")
    set_run_font(r1, bold=True)
    r2 = p.add_run(val or "")
    set_run_font(r2, bold=False)


def add_multiline_paragraphs(doc, text, size=11):
    if not text:
        return
    chunks = [c.strip() for c in re.split(r"\n\s*\n", text) if c.strip()]
    if not chunks:
        chunks = [text.strip()]
    for chunk in chunks:
        add_paragraph(doc, chunk, size=size)


# ===========================================
#          LOGO EN ENCABEZADO
# ===========================================

def insert_logo_header(doc, logo_path):
    if not logo_path or not Path(logo_path).exists():
        return
    for section in doc.sections:
        header = section.header
        if not header.paragraphs: p = header.add_paragraph()
        else:
            p = header.paragraphs[0]
            for run in p.runs: run.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run()
        run.add_picture(logo_path, width=Inches(4.2))


# ===========================================
#          RENDER DOCX
# ===========================================

def render_docx(out_path: Path, data: dict, logo_path: str):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.header_distance = Cm(0)

    insert_logo_header(doc, logo_path)

    add_paragraph(doc, "INFORME DE PRESENTACIÓN DE CANDIDATO/A", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=12)

    fecha = data.get("fecha") or datetime.now().strftime("%Y-%m-%d")
    add_label_value(doc, "Fecha", fecha)
    add_label_value(doc, "Puesto", data.get("puesto", ""))
    doc.add_paragraph()

    add_paragraph(doc, "Resumen Ejecutivo", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, before=6)
    add_paragraph(doc, data.get("resumen_ejecutivo", ""), size=11)

    add_paragraph(doc, "1.  Ficha Técnica", bold=True, size=12, before=12)
    ft = data.get("ficha_tecnica", {})
    
    table = doc.add_table(rows=0, cols=2)
    set_table_borders(table)
    table.autofit = False
    col_width = Inches(2.8)
    for campo, val in [
        ("Nombre", ft.get("nombre", "")),
        ("Ubicación", ft.get("ubicacion", "")),
        ("Nivel de experiencia", ft.get("nivel_experiencia", "")),
        ("Formación Formal", ft.get("formacion_formal", "")),
        ("Nivel Inglés", ft.get("nivel_ingles", "")),
        ("Disponibilidad", ft.get("disponibilidad", "")),
    ]:
        row = table.add_row().cells
        row[0].width = col_width; row[1].width = col_width
        r1 = row[0].paragraphs[0].add_run(campo); set_run_font(r1, bold=True)
        r2 = row[1].paragraphs[0].add_run(val); set_run_font(r2)
    doc.add_paragraph()

    # 2. Competencias Técnicas
    comps = data.get("competencias_tecnicas", [])
    if comps:
        add_paragraph(doc, "2.  Competencias Técnicas y destacada", bold=True, size=12, before=12)
        t = doc.add_table(rows=1, cols=2)
        set_table_borders(t)
        t.autofit = False
        t.rows[0].cells[0].width = col_width; t.rows[0].cells[1].width = col_width
        t.rows[0].cells[0].text = "Competencia"; t.rows[0].cells[1].text = "Nivel"
        for cell in t.rows[0].cells: set_run_font(cell.paragraphs[0].runs[0], bold=True)
        for item in comps:
            row = t.add_row().cells
            row[0].width = col_width; row[1].width = col_width
            set_run_font(row[0].paragraphs[0].add_run(item["competencia"]))
            set_run_font(row[1].paragraphs[0].add_run(item["nivel"]))
        doc.add_paragraph()

    # 2.1 Habilidades Blandas
    soft = data.get("habilidades_blandas", [])
    if soft:
        add_paragraph(doc, "2.1  Habilidades Blandas (Soft Skills)", bold=True, size=12, before=12)
        t_soft = doc.add_table(rows=1, cols=2)
        set_table_borders(t_soft)
        t_soft.autofit = False
        t_soft.rows[0].cells[0].width = col_width; t_soft.rows[0].cells[1].width = col_width
        t_soft.rows[0].cells[0].text = "Habilidad"; t_soft.rows[0].cells[1].text = "Nivel"
        for cell in t_soft.rows[0].cells: set_run_font(cell.paragraphs[0].runs[0], bold=True)
        for item in soft:
            row = t_soft.add_row().cells
            row[0].width = col_width; row[1].width = col_width
            set_run_font(row[0].paragraphs[0].add_run(item["habilidad"]))
            set_run_font(row[1].paragraphs[0].add_run(item["nivel"]))
        doc.add_paragraph()

    # 3. Herramientas
    tools = data.get("herramientas", [])
    if tools:
        add_paragraph(doc, "3.  Herramientas y Software", bold=True, size=12, before=12)
        t2 = doc.add_table(rows=1, cols=2)
        set_table_borders(t2)
        t2.autofit = False
        t2.rows[0].cells[0].width = col_width; t2.rows[0].cells[1].width = col_width
        t2.rows[0].cells[0].text = "Herramienta"; t2.rows[0].cells[1].text = "Nivel de Dominio"
        for cell in t2.rows[0].cells: set_run_font(cell.paragraphs[0].runs[0], bold=True)
        for item in tools:
            row = t2.add_row().cells
            row[0].width = col_width; row[1].width = col_width
            set_run_font(row[0].paragraphs[0].add_run(item["herramienta"]))
            set_run_font(row[1].paragraphs[0].add_run(item["nivel"]))
        doc.add_paragraph()

    plus = data.get("plus", "").strip()
    if plus:
        add_paragraph(doc, "4.  Plus", bold=True, size=12, before=12)
        add_multiline_paragraphs(doc, plus, size=11)

    fs = data.get("formacion_sugerida", "").strip()
    if fs:
        add_paragraph(doc, "5.  Formación Sugerida", bold=True, size=12, before=12)
        add_paragraph(doc, fs, size=11)

    rf = data.get("recomendacion_final", "").strip()
    if rf:
        add_paragraph(doc, "6.  Recomendación Final", bold=True, size=12, before=12)
        add_multiline_paragraphs(doc, rf, size=11)

    resp = data.get("responsable", "")
    if resp:
        p = doc.add_paragraph()
        r1 = p.add_run("Responsable: "); set_run_font(r1, bold=True)
        r2 = p.add_run(resp); set_run_font(r2)

    notas_input = data.get("_meta", {}).get("notas_raw", "").strip()
    if notas_input:
        doc.add_page_break()
        add_paragraph(doc, "Notas Adicionales / Input del Reclutador", bold=True, size=12, before=12)
        add_multiline_paragraphs(doc, notas_input, size=11)

    doc.save(str(out_path))


# ===========================================
#                 MAIN
# ===========================================

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--extra_file", type=str, default="")
    ap.add_argument("--extra", type=str, default="")
    ap.add_argument("--cv", type=str, default="")
    ap.add_argument("--outdir", type=str, required=True)
    ap.add_argument("--basename", type=str, required=True)
    ap.add_argument("--logo", type=str, default="")
    ap.add_argument("--no-pdf", action="store_true")
    args = ap.parse_args()

    full_text = args.extra or ""
    if args.extra_file:
        full_text = read_text(Path(args.extra_file)) or full_text

    cv_text = read_cv_any(Path(args.cv)) if args.cv else ""
    
    # 1. Intentamos extraer herramientas con %
    tools = extract_tools_with_percent(full_text)
    
    # 2. IA (Con prompt mejorado para extraer si no hubo %)
    data = ai_extract_fields(full_text, cv_text, tools)

    # 3. Normalizador (Por si la IA se confundió)
    data = normalize_skills(data)

    if "_meta" not in data: data["_meta"] = {}
    data["_meta"]["notas_raw"] = full_text

    outdir = Path(args.outdir)
    outdir.mkdir(parents=True, exist_ok=True)
    out_path = outdir / f"{args.basename}.docx"
    render_docx(out_path, data, args.logo)

    print(json.dumps({"ok": True, "docx": str(out_path), "mapping": data}, ensure_ascii=False))


if __name__ == "__main__":
    main()