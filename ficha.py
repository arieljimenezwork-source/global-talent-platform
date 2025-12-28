# -*- coding: utf-8 -*-
"""
ficha.py — Genera Ficha de Evaluación (DOCX/PDF) a partir de TXT + CV con un solo llamado a IA.
- La IA devuelve TODOS los campos (encabezado, info personal, resumen, áreas técnicas,
  blandas con justificación, herramientas con %, compatibilidad, potencial, formativas,
  recomendación final).
- Herramientas: SOLO desde el TXT con % (regex). La IA no puede inventarlas.
- Profesión: prioriza lo que esté en el CV.
- Contabilidad: criterio de compatibilidad más estricto.
- Secciones/filas vacías NO se muestran.
- Validador PERMISIVO: no corta la ejecución; solo advierte en stderr.

🔸 Personalización pedida:
Si en el texto (extra_file / extra) aparece una línea:
  "Tipo asistente: <texto libre>"  o  "Tipo de asistente: <texto libre>"
ese valor se usa literalmente en el título:
  Ficha de Evaluación de Perfil – Asistente virtual de "<texto libre>"
"""

import os, re, json, argparse, subprocess
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Tuple, Any

# --- Forzar stdout/stderr a UTF-8 en Windows ---
import sys, io
try:
    # Python 3.7+ soporta reconfigure
    sys.stdout.reconfigure(encoding="utf-8", errors="strict")
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    # Fallback por si el stream no permite reconfigure
    if hasattr(sys.stdout, "buffer"):
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="strict")
    if hasattr(sys.stderr, "buffer"):
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")


# === Opcional: cargar .env ===
try:
    from dotenv import load_dotenv
    load_dotenv()
except Exception:
    pass

# === OpenAI ===
_HAS_OPENAI = True
try:
    from openai import OpenAI
except Exception:
    _HAS_OPENAI = False

# === Lectura CV/TXT y DOCX ===
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

try:
    from docx2pdf import convert as docx2pdf_convert
except Exception:
    docx2pdf_convert = None

try:
    from pdfminer.high_level import extract_text as pdf_extract_text
    PDFMINER_AVAILABLE = True
except Exception:
    PDFMINER_AVAILABLE = False

try:
    import docx2txt
    DOCTXTXT_AVAILABLE = True
except Exception:
    DOCTXTXT_AVAILABLE = False


# =========================
#   Config
# =========================
PLACE = {"areas": 10, "competencias": 5, "herrs": 12, "pots": 5, "recs": 5}
_CATEGORIES = ["Contabilidad","Administración","IA","Recursos humanos","Diseño",
               "Marketing","Economía","Comunicación","Calidad","Arquitectura"]

# Usos conocidos (ampliado)
TOOLS_USAGE = {
    "autocad": "Diseño y planos arquitectónicos",
    "revit": "Modelado BIM y coordinación",
    "sketchup": "Modelado 3D y visualización",
    "archicad": "Modelado BIM y documentación",
    "photoshop": "Edición y retoque de imágenes",
    "illustrator": "Gráficos vectoriales y piezas",
    "microsoft project": "Planificación y seguimiento",
    "microsoft office": "Documentación y presentaciones",
    "microsoft excel": "Modelado y análisis de datos en hojas de cálculo",
    "excel": "Modelado y análisis de datos en hojas de cálculo",
    "google workspace": "Colaboración y productividad",
    "google calendar": "Gestión de agenda, eventos y reuniones",
    "calendar": "Gestión de agenda, eventos y reuniones",
    "asana": "Gestión y priorización de tareas",
    "trello": "Tableros kanban y seguimiento",
    "notion": "Documentación de procesos",
    "metricool": "Analítica y programación social",
    "mailchimp": "Email marketing y automatización",
    "pipedrive": "Gestión de ventas y CRM",
    "canva": "Diseño rápido de piezas",
    "meta ads": "Gestión de campañas en Meta",
    "google ads": "Publicidad y performance",
    "odoo": "ERP y procesos contables",
    "sap": "ERP corporativo",
    "profit": "Contabilidad y finanzas",
    "slack": "Mensajería y colaboración en equipos",
    "discord": "Comunicación por canales y comunidades",
}

# Colores y fuente
COLOR_BANNER = "0B2A5B"
COLOR_HDR = "E7EEF8"
COLOR_SECTION_BAR = "F2F2F2"
FONT_NAME = "Arial"
FS_BASE = 10
FS_TITLE = 11

SIDE_PADDING_CM = 1.3
W_TECH = [2.5, 1.2]
W_SOFT = [2.2, 5.5]
W_TOOLS = [3, 4, 1.5]
W_COMPAT = [3, 2, 4]


# =========================
#   Utilidades
# =========================
def read_text(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="ignore")
    except Exception:
        return ""

def read_cv_any(path: Path) -> str:
    if not path or not path.exists():
        return ""
    suf = path.suffix.lower()
    try:
        if suf == ".docx" and DOCTXTXT_AVAILABLE:
            return docx2txt.process(str(path)) or ""
        if suf == ".pdf" and PDFMINER_AVAILABLE:
            return pdf_extract_text(str(path)) or ""
        return read_text(path)
    except Exception:
        return ""

def sanitize_short_field(s: str, max_chars=48, max_words=6) -> str:
    s = (s or "").strip()
    s = re.sub(r"\s+", " ", s)
    words = re.findall(r"\S+", s)
    if len(s) > max_chars or len(words) > max_words:
        return ""
    return s

def pct_int(s: str) -> int:
    m = re.search(r"(\d{1,3})", (s or ""))
    return int(m.group(1)) if m else -1

def normalize_tool(name: str) -> str:
    """
    Limpia frases tipo 'He utilizado slack', 'uso de Calendar', etc. y
    normaliza sinónimos a la marca correcta.
    """
    s = (name or "").strip()
    s = re.sub(r"^[^A-Za-zÁ-ÿ]+", "", s)
    s = re.sub(
        r"(?i)\b(he\s+utilizado|he\s+usado|he\s+manejado|utilizo|uso|manejo|maneja|manejar|"
        r"conocimiento\s+de|experiencia\s+en|trabajo\s+con|he\s+trabajado\s+con|uso\s+de|"
        r"herramientas?\s+de)\b[:\-]?\s*",
        "",
        s,
    )
    s = re.sub(r"\s+", " ", s).strip()

    sl = s.lower()
    syn = {
        "project": "Microsoft Project",
        "ms project": "Microsoft Project",
        "office": "Microsoft Office",
        "excel": "Microsoft Excel",
        "g suite": "Google Workspace",
        "google suite": "Google Workspace",
        "calendar": "Google Calendar",
        "google calendar": "Google Calendar",
        "autocad": "AutoCAD",
        "sketchup": "SketchUp",
        "adwords": "Google Ads",
        "facebook ads": "Meta Ads",
        "slack": "Slack",
        "discord": "Discord",
    }
    norm = syn.get(sl, s)
    if norm and norm[0].islower():
        norm = norm[0].upper() + norm[1:]
    return norm

def tool_usage_guess(name: str) -> str:
    """
    Devuelve un uso claro y específico; nunca 'Uso relacionado al rol'.
    """
    key = (name or "").strip().lower()
    if key in TOOLS_USAGE:
        return TOOLS_USAGE[key]

    # Heurísticas por familia
    if "excel" in key or "sheet" in key:
        return "Modelado y análisis de datos en hojas de cálculo"
    if "calendar" in key or "agenda" in key:
        return "Gestión de agenda, eventos y reuniones"
    if any(w in key for w in ["slack", "discord", "teams", "whatsapp", "telegram"]):
        return "Mensajería y colaboración en equipos"
    if any(w in key for w in ["ads", "sem", "seo", "mailchimp", "metricool"]):
        return "Gestión y optimización de campañas de marketing digital"
    if any(w in key for w in ["crm", "pipedrive", "hubspot", "salesforce"]):
        return "Gestión de oportunidades comerciales en CRM"
    if any(w in key for w in ["project", "asana", "trello", "notion", "jira"]):
        return "Planificación y seguimiento de proyectos y tareas"
    if any(w in key for w in ["autocad", "revit", "archicad", "sketchup"]):
        return "Modelado y documentación técnica"
    if any(w in key for w in ["photoshop", "illustrator", "canva", "figma"]):
        return "Diseño y edición de piezas visuales"
    if any(w in key for w in ["odoo", "sap", "profit"]):
        return "Soporte a procesos contables y de gestión en ERP"

    return "Aplicación práctica de herramientas del área"

def infer_years_experience(main_txt: str, cv_txt: str) -> int:
    # explícito
    def _explicit_years(t: str) -> int:
        t = (t or "").lower()
        cands = []
        for m in re.finditer(r"(?:con\s+)?(\d{1,2})\s+años\s+de\s+experiencia", t): cands.append(int(m.group(1)))
        for m in re.finditer(r"experiencia\s+de\s+(\d{1,2})\s+años", t): cands.append(int(m.group(1)))
        for m in re.finditer(r"(?:m[aá]s\s+de|over)\s+(\d{1,2})\s+años", t): cands.append(int(m.group(1))+1)
        return max(cands) if cands else -1
    for source in [main_txt, cv_txt]:
        y = _explicit_years(source)
        if y >= 0:
            return y

    # inferencia por rango de años en experiencia
    text = (cv_txt or "") + "\n" + (main_txt or "")
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    YEAR_RE  = re.compile(r"\b(19\d{2}|20\d{2})\b")
    IN_EXP_RE = re.compile(r"(?i)(experiencia|laboral|trayectoria|work\s*experience|employment)")
    BAD_NEAR = re.compile(r"(?i)(nacim|nacimiento|edad|fecha|dd|mm|aaaa)")
    in_exp = False
    years: List[int] = []
    curr = datetime.now().year
    for ln in lines:
        if IN_EXP_RE.search(ln): in_exp = True; continue
        if re.search(r"(?i)(educaci[oó]n|formaci[oó]n|estudios|referencias|datos|perfil|sobre m[ií])", ln):
            in_exp = False
        if not in_exp: continue
        if BAD_NEAR.search(ln): continue
        for m in YEAR_RE.findall(ln):
            y = int(m)
            if 1980 <= y <= curr:
                years.append(y)
    if years:
        start = min(years)
        est = max(0, curr - start)
        if est <= 45: return est
    return 0


# =========================
#   Herramientas (solo con %)
# =========================
TOOLS_ANCHOR_RE = re.compile(
    r"(?is)para\s+el\s+área\s+que\s+est[áa]s?\s+aplicando.*?(herramientas|programas).*?(utilizado|utilizas|has utilizado)"
)
NEXT_SECTION_RE = re.compile(
    r"(?im)^\s*(secci[oó]n|compatibilidad|dudas|inquietudes|disponibilidad|expectativas|observaci[oó]n|categor[ií]a|notas)\b.*$"
)

def _slice_tools_block(full_text: str) -> str:
    s = full_text.replace("\r", "")
    m = TOOLS_ANCHOR_RE.search(s)
    if not m:
        return ""
    start = m.start()
    tail = s[start:]
    end_match = NEXT_SECTION_RE.search(tail)
    return tail[:end_match.start()] if end_match else tail

def extract_tools_with_percent_from_text(full_text: str, maxn: int = 40) -> List[Dict[str, str]]:
    """
    Extrae SOLO herramientas con % explícito desde el texto.
    Soporta:
      - "Asana – 100%"  /  "Asana - 100%"  /  "Asana 100%"
      - Nombre en una línea y % en la siguiente.
    """
    text = (full_text or "").replace("\r", "")
    block = _slice_tools_block(text)
    candidates: List[Tuple[str, str]] = []

    def harvest(lines: List[str]) -> None:
        nonlocal candidates
        pat_inline = re.compile(r"([A-Za-zÁ-ÿ][A-Za-zÁ-ÿ0-9 .+/_()-]{1,}?)\s*[:\-–—]?\s*(\d{1,3})\s*%", re.I)
        pat_name = re.compile(r"([A-Za-zÁ-ÿ][A-Za-zÁ-ÿ0-9 .+/_()-]{1,})$")
        buffer = None
        for raw in lines:
            ln = raw.strip()
            if not ln: continue
            hits = pat_inline.findall(ln)
            if hits:
                for n, p in hits:
                    n = normalize_tool(n)
                    if len(re.sub(r"[^A-Za-zÁ-ÿ]", "", n)) >= 3:
                        candidates.append((n, f"{min(max(int(p),0),100)}%"))
                buffer = None
                continue
            if buffer is None:
                m = pat_name.search(ln)
                if m:
                    cand = normalize_tool(m.group(1))
                    if len(re.sub(r"[^A-Za-zÁ-ÿ]", "", cand)) >= 3:
                        buffer = cand
                continue
            else:
                m2 = re.search(r"(\d{1,3})\s*%", ln)
                if m2:
                    candidates.append((buffer, f"{min(max(int(m2.group(1)),0),100)}%"))
                    buffer = None

    if block:
        harvest(block.splitlines())
    if not candidates:
        harvest(text.splitlines())

    merged: Dict[str, Dict[str, str]] = {}
    for name, dom in candidates:
        key = name.lower()
        if key in merged:
            if pct_int(dom) > pct_int(merged[key]["dom"]):
                merged[key]["dom"] = dom
        else:
            merged[key] = {"herr": name, "dom": dom, "uso": tool_usage_guess(name)}
    out = list(merged.values())
    out.sort(key=lambda x: pct_int(x.get("dom","")), reverse=True)
    return out[:maxn]


# =========================
#   IA: un solo llamado para TODO
# =========================
def _ai_enabled() -> bool:
    return _HAS_OPENAI and bool(os.environ.get("OPENAI_API_KEY"))

def ai_extract_all_fields(full_text: str, cv_text: str,
                          pre_tools: List[Dict[str, str]],
                          model: str = "gpt-4o-mini",
                          temperature: float = 0.2) -> Dict[str, Any]:
    """
    Un solo llamado a la IA con TXT + CV. Devuelve todos los campos para la ficha.
    Herramientas se pasan desde regex (pre_tools) y la IA debe usarlas tal cual.
    """
    if not _ai_enabled():
        raise RuntimeError("OPENAI_API_KEY no configurada o SDK no disponible.")

    client = OpenAI()
    fuente = ((full_text or "") + "\n\n---CV---\n\n" + (cv_text or "")).strip()
    years_rules = infer_years_experience(full_text, cv_text)

    # Cantidades como GUÍA (no bloqueante). La validación es permisiva.
    system = (
        "Eres un analista de RR.HH. Devuelve EXCLUSIVAMENTE un JSON válido UTF-8, "
        "SIN texto extra ni fences, siguiendo este esquema. "
        "No inventes datos: deben salir del TXT o del CV. Profesión prioriza CV. País literal. "
        f"Área del título debe ser UNA de: {', '.join(_CATEGORIES)}. "
        "Edad: infiérela de la fecha de nacimiento si no está explícita. "
        "Resumen ejecutivo (guía 52–56 palabras): tercera persona, tono corporativo; "
        "menciona años de experiencia; si no están explícitos, infiérelos y decláralo. "
        "Competencias técnicas (guía 6–10): no blandas ni herramientas, puedes inferir las competencias tecnicas clave de las herrmientas ydel contexto general  para completar las 6 competencias minimas; cada una con evaluación Alta/Media/Básica. "
        "Siempre listar mínimo 6 competencias, sin repeticiones ni sinónimos"
        "NO incluir competencias blandas, actitudes ni nombres de herramientas. "
        "Cada competencia debe redactarse en formato genérico, técnico y conciso. "
        "Competencias blandas (guía 3–5): cada una con justificación ≤25 palabras. "
        "HERRAMIENTAS (estricto): Usa EXACTAMENTE la lista recibida en 'tools_regex' "
        "(misma cantidad y mismo 'dom' en %). Para cada elemento: "
        "1) Normaliza el nombre para que sea únicamente el software/plataforma/aplicación (p.ej., "
        "'He utilizado slack' -> 'Slack', 'excel' -> 'Microsoft Excel'). "
        "2) PROHIBIDO usar frases, verbos u oraciones como nombre; solo el nombre real de la herramienta. "
        "3) Escribe 'uso' como una frase breve y concreta del uso GENERAL de la herramienta "
        "(p.ej., 'Modelado y análisis de datos en hojas de cálculo' para Microsoft Excel; "
        "'Gestión de agenda, eventos y reuniones' para Google Calendar; "
        "'Comunicación por canales y comunidades' para Discord; "
        "'Mensajería y colaboración en equipos' para Slack). "
        "4) No uses descripciones genéricas como 'Uso relacionado al rol'. "
        "Para cada herramienta incluye 'uso' (descripción breve del uso general) y 'dom' (NN%). "
        "Compatibilidad: cuatro categorías (Marketing Digital, Ventas y E-commerce, Gestión y Control, Contabilidad). "
        "Califica 1–5 estrellas con criterios estrictos basados en evidencia textual; incluye 'obs' breve y 'evid' con citas literales. "
        "Potencial: cinco bullets.describiendo el ptencial de aporte auna empresa que presenta el aplicante"
        "Recomendaciones Formativas: hasta cinco bullets.mencionar areas de mejora o que debe fortalecer tanto tecnicas como blandas para que el aplicante mejore como asistente "
        "Recomendación final (guía 56–60 palabras). "
        "Aunque alguna sección tenga poca evidencia, SIEMPRE devuelve el esquema completo."
    )

    expected = {
        "area_titulo": "string (una de la lista; p.ej., Arquitectura)",
        "info_personal": {
            "nombre": "string",
            "edad": "string",
            "region": "string (país literal)",
            "profesion": "string (prioriza CV; vacío si no hay)",
            "categoria": "string (igual a area_titulo, si aplica)"
        },
        "resumen": "string",
        "tech_areas": [{"area": "string", "eval": "Alta|Media|Básica"}],
        "soft_skills": [{"comp":"string", "just":"string <=25 palabras"}],
        "tools": [{"herr":"string","uso":"string breve","dom":"NN%"}],  # EXACTAS a pre_tools
        "compat": {
            "Marketing Digital":{"rating":1,"obs":"string","evid":["frag1","frag2"]},
            "Ventas y E-commerce":{"rating":1,"obs":"string","evid":["frag1","frag2"]},
            "Gestión y Control":{"rating":1,"obs":"string","evid":["frag1","frag2"]},
            "Contabilidad":{"rating":1,"obs":"string","evid":["frag1","frag2"]},
        },
        "potencial": ["bullet1","bullet2","bullet3","bullet4","bullet5"],
        "formativas": ["bullet","bullet","bullet","bullet","bullet (máx)"],
        "recomendacion_final": "string"
    }

    user = {
        "texto_fuente": fuente,
        "años_experiencia_reglas": years_rules,
        "categorias_validas": _CATEGORIES,
        "tools_regex": pre_tools,  # La IA debe usarlas tal cual
        "schema": expected
    }

    resp = client.responses.create(
        model=model,
        temperature=temperature,
        input=[
            {"role":"system","content": system},
            {"role":"user","content": json.dumps(user, ensure_ascii=False)}
        ]
    )

    raw = (resp.output_text or "").strip()
    raw = re.sub(r"^```(json)?\s*|\s*```$", "", raw).strip()
    data = json.loads(raw)

    # Forzar herramientas desde regex (mantiene 'uso' si la IA lo dio; si no, se infiere)
    fixed = []
    usage_map = { (t.get("herr","").strip().lower(), t.get("dom","").strip()): None for t in pre_tools }
    for t in (data.get("tools") or []):
        key = ((t.get("herr","").strip().lower()), (t.get("dom","").strip()))
        if key in usage_map:
            usage_map[key] = (t.get("uso") or "").strip() or tool_usage_guess(t.get("herr"))
    for pt in pre_tools:
        k = (pt["herr"].strip().lower(), pt["dom"].strip())
        use = usage_map.get(k) or tool_usage_guess(pt["herr"])
        fixed.append({"herr": pt["herr"], "dom": pt["dom"], "uso": use})
    data["tools"] = fixed

    return data


# =========================
#   Validador PERMISIVO (no aborta)
# =========================
def _soft_validate_ai_payload(data: dict, pre_tools: List[Dict[str, str]]) -> list[str]:
    warnings: list[str] = []

    def w(msg: str):
        clean = (msg or "").replace("\u2013", "-").replace("\u2014", "-").replace("\u2019","'")
        warnings.append(clean)

    # Estructura básica
    required_keys = ["area_titulo","info_personal","resumen","tech_areas","soft_skills",
                     "tools","compat","potencial","formativas","recomendacion_final"]
    for k in required_keys:
        if k not in data:
            w(f"falta clave '{k}' en el JSON de la IA")

    ip = (data.get("info_personal") or {})
    for k in ["nombre","edad","region","profesion","categoria"]:
        if k not in ip:
            w(f"info_personal.{k} ausente")

    # Herramientas: coherencia con regex (nombre + %)
    tools = data.get("tools") or []
    def _key(t): return ((t.get("herr","").strip().lower()), (t.get("dom","").strip()))
    pre_set = {_key(t) for t in (pre_tools or [])}
    got_set = {_key(t) for t in tools}
    if pre_set != got_set:
        w("tools de la IA no coincide exactamente con tools_regex (nombre y %). Se forzarán desde regex en post-proceso.")

    # Compat: presencia de 4 categorías y forma general
    compat = data.get("compat") or {}
    for cat in ["Marketing Digital","Ventas y E-commerce","Gestión y Control","Contabilidad"]:
        if cat not in compat:
            w(f"compat.{cat} ausente")
        else:
            c = compat.get(cat) or {}
            if "rating" not in c: w(f"compat.{cat}.rating ausente")
            if "obs" not in c:    w(f"compat.{cat}.obs ausente")
            if "evid" not in c or not isinstance(c.get("evid"), list):
                w(f"compat.{cat}.evid ausente o no es lista")

    # Vacíos: solo aviso
    if not (data.get("resumen") or "").strip():
        warnings.append("resumen vacío")
    if not (data.get("recomendacion_final") or "").strip():
        warnings.append("recomendacion_final vacía")
    if not (data.get("tech_areas") or []):
        warnings.append("tech_areas viene vacío")
    if not (data.get("soft_skills") or []):
        warnings.append("soft_skills viene vacío")
    if not (data.get("potencial") or []):
        warnings.append("potencial viene vacío")
    if not (data.get("formativas") or []):
        warnings.append("formativas viene vacío")

    return warnings


# =========================
#   Render helpers (DOCX)
# =========================
def _clean_val(x: str) -> str:
    return (x or "").replace("\u00A0", " ").strip()

def _apply_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def _set_run_font(run, bold=False, size_pt=FS_BASE, color_hex=None):
    run.bold = bold
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)

def _set_paragraph_compact(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.05
    if align is not None:
        p.alignment = align

def _set_table_borders(tbl, size=8, color="000000"):
    tbl_el = tbl._tbl
    tblPr = tbl_el.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl_el._element.append(tblPr)
    borders = OxmlElement('w:tblBorders')
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn('w:val'), 'single' if size else 'nil')
        element.set(qn('w:sz'), str(size))
        element.set(qn('w:color'), color)
        borders.append(element)
    for child in list(tblPr):
        if child.tag == qn('w:tblBorders'):
            tblPr.remove(child)
    tblPr.append(borders)
    tbl.autofit = False

def _content_width_cm(doc: Document) -> float:
    sec = doc.sections[0]
    width_cm = (sec.page_width - sec.left_margin - sec.right_margin) / 360000.0
    return float(width_cm)

def _fix_table_width(doc: Document, tbl, weights: List[float]):
    total_cm = max(1.0, _content_width_cm(doc) - 2*SIDE_PADDING_CM)
    s = sum(weights) if sum(weights) > 0 else 1.0
    widths_cm = [total_cm * (w / s) for w in weights]
    tbl.autofit = False
    for j, w in enumerate(widths_cm):
        for row in tbl.rows:
            try:
                row.cells[j].width = Cm(w)
            except Exception:
                pass
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

def _set_table_fullwidth_1col(doc: Document, tbl):
    _fix_table_width(doc, tbl, [1])

def _add_bar(doc: Document, text: str, color_hex: str, font_color="000000"):
    tbl = doc.add_table(rows=1, cols=1)
    _set_table_borders(tbl, size=8, color="000000")
    cell = tbl.rows[0].cells[0]
    _apply_cell_shading(cell, color_hex)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_compact(p, align=None)
    run = p.add_run(text)
    _set_run_font(run, bold=True, size_pt=FS_TITLE, color_hex=font_color)
    _set_table_fullwidth_1col(doc, tbl)
    return tbl

def _add_section_bar(doc: Document, text: str):
    return _add_bar(doc, text, COLOR_SECTION_BAR)

def stars(n: int) -> str:
    n = max(1, min(5, int(n or 1)))
    return "★" * n + "☆" * (5 - n)

def _maybe_add_tech(doc: Document, tech_rows: List[Tuple[str,str]]):
    rows = [(a or "", e or "") for a, e in tech_rows if (a or e)]
    if not rows:
        return
    _add_section_bar(doc, "Competencias Técnicas Clave")
    tbl = doc.add_table(rows=1, cols=2); _set_table_borders(tbl, 8)
    hdr = tbl.rows[0].cells; hdr[0].text = "Área"; hdr[1].text = "Evaluación"
    for c in hdr:
        _apply_cell_shading(c, COLOR_HDR)
        if c.paragraphs[0].runs: _set_run_font(c.paragraphs[0].runs[0], bold=True)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for a, e in rows:
        row = tbl.add_row().cells
        row[0].text, row[1].text = a or "—", e or "—"
    _fix_table_width(doc, tbl, W_TECH)

def _maybe_add_soft(doc: Document, soft_rows: List[Tuple[str,str]]):
    rows = [(c or "", j or "") for c, j in soft_rows if (c or j)]
    if not rows:
        return
    _add_section_bar(doc, "Competencias Destacadas")
    tbl = doc.add_table(rows=1, cols=2); _set_table_borders(tbl, 8)
    hdr = tbl.rows[0].cells; hdr[0].text = "Competencia"; hdr[1].text = "Justificación"
    for c in hdr:
        _apply_cell_shading(c, COLOR_HDR)
        if c.paragraphs[0].runs: _set_run_font(c.paragraphs[0].runs[0], bold=True)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for c, j in rows:
        row = tbl.add_row().cells
        row[0].text, row[1].text = c or "—", j or "—"
    _fix_table_width(doc, tbl, W_SOFT)

def _maybe_add_tools(doc: Document, tools: List[Dict[str,str]]):
    rows = []
    for t in tools:
        h = _clean_val(t.get("herr",""))
        u = _clean_val(t.get("uso",""))
        d = _clean_val(t.get("dom",""))
        if h and d:  # requiere %; si falta uso, igual se muestra pero intentamos siempre ponerlo
            rows.append([h, u or "—", d])
    if not rows:
        return
    _add_section_bar(doc, "Herramientas")
    tbl = doc.add_table(rows=1, cols=3); _set_table_borders(tbl, 8)
    hdr = tbl.rows[0].cells; hdr[0].text, hdr[1].text, hdr[2].text = "Herramienta", "Uso", "Dominio (%)"
    for c in hdr:
        _apply_cell_shading(c, COLOR_HDR)
        if c.paragraphs[0].runs: _set_run_font(c.paragraphs[0].runs[0], bold=True)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in rows:
        row = tbl.add_row().cells
        row[0].text, row[1].text, row[2].text = r[0], r[1], r[2]
    _fix_table_width(doc, tbl, W_TOOLS)

def _is_tool_like(text: str, tool_names_lower: set[str]) -> bool:
    s = (text or "").strip().lower()
    if not s:
        return False
    if any(t in s for t in tool_names_lower):
        return True
    if re.search(r"\b\d{1,3}\s*%|\b(v|versi[oó]n)\s*\d+", s):
        return True
    if re.search(r"\b(autocad|revit|sketchup|archicad|photoshop|illustrator|odoo|sap|profit|excel|word|canva|mailchimp|metricool|pipedrive|notion|asana|trello|project|google ads|meta ads)\b", s):
        return True
    return False

def _normalize_eval_from_pct(pct: int) -> str:
    try:
        v = int(pct)
    except Exception:
        v = -1
    if v >= 80: return "Alta"
    if v >= 50: return "Media"
    return "Básica" if v >= 0 else "Básica"

def _enforce_tech_areas_rules(data: dict) -> None:
    """Solo limpia técnicas inválidas; NO agrega placeholders."""
    tools = { (t.get("herr","").strip().lower()) for t in (data.get("tools") or []) }
    filtered = []
    seen = set()
    for item in (data.get("tech_areas") or []):
        area = (item.get("area") or "").strip()
        ev = (item.get("eval") or "").strip().capitalize()
        if not area or ev not in {"Alta","Media","Básica"}:
            continue
        if _is_tool_like(area, tools):
            continue
        key = re.sub(r"\s+", " ", area.lower())
        if key in seen: 
            continue
        seen.add(key)
        filtered.append({"area": area, "eval": ev})
    data["tech_areas"] = filtered

def _justify_document(doc: Document) -> None:
    def _justify_para(p):
        if p.alignment in (WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT):
            return
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for p in doc.paragraphs:
        for r in p.runs:
            if r.font.size is None: r.font.size = Pt(FS_BASE)
            if r.font.name is None: r.font.name = FONT_NAME
        _justify_para(p)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _justify_para(p)

def render_docx(out_docx: Path, data: Dict[str, Any], logo_path: str | None = None):
    doc = Document()

    # Logo
    if logo_path and Path(logo_path).exists():
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture(logo_path, width=Cm(11.0))
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Tabla inicial (encabezado, info personal, resumen)
    tbl = doc.add_table(rows=3, cols=1)
    _set_table_borders(tbl, 8, "000000")

    # ==== TÍTULO PRINCIPAL (centrado con espacio adecuado) ====
    titulo_usuario = _clean_val(data.get("titulo_usuario", ""))
    area_ia = _clean_val(data.get("area_titulo", ""))
    if titulo_usuario:
        titulo_banner = f"Ficha de Evaluación de Perfil – Asistente virtual de {titulo_usuario}"
    elif area_ia:
        titulo_banner = f"Ficha de Evaluación de Perfil – Asistente virtual de {area_ia}"
    else:
        titulo_banner = "Ficha de Evaluación de Perfil – Asistente virtual"

    cell_title = tbl.rows[0].cells[0]
    _apply_cell_shading(cell_title, COLOR_BANNER)
    p = cell_title.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_compact(p, align=None)
    run = p.add_run(titulo_banner)
    _set_run_font(run, bold=True, size_pt=FS_TITLE, color_hex="FFFFFF")
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(14)

    # ==== INFORMACIÓN PERSONAL ====
    info = data.get("info_personal") or {}
    cell_info = tbl.rows[1].cells[0]
    pinfo = cell_info.paragraphs[0]
    _set_paragraph_compact(pinfo)
    r = pinfo.add_run("Información Personal:")
    _set_run_font(r, bold=True, size_pt=FS_TITLE)

    def info_line(cell, label, value):
        val = _clean_val(value)
        if not val:
            return
        par = cell.add_paragraph()
        _set_paragraph_compact(par)
        r0 = par.add_run("• ")
        _set_run_font(r0, bold=True)
        r1 = par.add_run(f"{label}: ")
        _set_run_font(r1, bold=True)
        par.add_run(val)

    nombre_raw = info.get("nombre", "")
    if nombre_raw:
        info_line(cell_info, "Nombre", " ".join(nombre_raw.split()[:2]))
    info_line(cell_info, "Edad", info.get("edad", ""))
    info_line(cell_info, "Región", info.get("region", ""))
    info_line(cell_info, "Profesión", info.get("profesion", ""))

    # ==== RESUMEN EJECUTIVO ====
    cell_sum = tbl.rows[2].cells[0]
    p2 = cell_sum.paragraphs[0]
    _set_paragraph_compact(p2)
    r2 = p2.add_run("Resumen Ejecutivo: ")
    _set_run_font(r2, bold=True)
    resumen = _clean_val(data.get("resumen", "")) or "—"
    p2.add_run(resumen)
    _set_table_fullwidth_1col(doc, tbl)

    # ==== COMPETENCIAS TÉCNICAS CLAVE ====
    tech_rows = [
        (_clean_val(itm.get("area", "")), _clean_val(itm.get("eval", "")))
        for itm in (data.get("tech_areas") or [])[:PLACE["areas"]]
        if _clean_val(itm.get("area", "")) or _clean_val(itm.get("eval", ""))
    ]
    if tech_rows:
        _maybe_add_tech(doc, tech_rows)

    # ==== COMPETENCIAS DESTACADAS ====
    soft_rows = [
        (_clean_val(itm.get("comp", "")), _clean_val(itm.get("just", "")))
        for itm in (data.get("soft_skills") or [])[:PLACE["competencias"]]
        if _clean_val(itm.get("comp", "")) or _clean_val(itm.get("just", ""))
    ]
    if soft_rows:
        _maybe_add_soft(doc, soft_rows)

    # ==== HERRAMIENTAS TÉCNICAS Y TECNOLOGÍAS ====
    if data.get("tools"):
        original_add_bar = _add_section_bar
        def _custom_add_section_bar(doc_local, text_local):
            return original_add_bar(doc_local, "Herramientas Técnicas y Tecnologías")
        globals()['_add_section_bar'] = _custom_add_section_bar
        _maybe_add_tools(doc, data.get("tools") or [])
        globals()['_add_section_bar'] = original_add_bar

    # ==== COMPATIBILIDAD ====
    def _maybe_add_compat(doc: Document, compat: Dict[str, Dict[str, Any]]):
        if not compat:
            return
        rows = []
        def add(cat):
            if cat in compat and isinstance(compat[cat], dict):
                r = compat[cat].get("rating", 2)
                o = compat[cat].get("obs", "")
                rows.append([cat, stars(r), o or ("No aplica" if (r or 0) <= 2 else "—")])
        for cat in ["Marketing Digital", "Ventas y E-commerce", "Gestión y Control", "Contabilidad"]:
            add(cat)
        if not rows:
            return
        _add_section_bar(doc, "Compatibilidad con Categorías de Asistencia Virtual")
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tbl = doc.add_table(rows=1, cols=3)
        _set_table_borders(tbl, 8)
        hdr = tbl.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Categoría", "Calificación", "Observación"
        for c in hdr:
            _apply_cell_shading(c, COLOR_HDR)
            if c.paragraphs[0].runs:
                _set_run_font(c.paragraphs[0].runs[0], bold=True)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in rows:
            row = tbl.add_row().cells
            row[0].text, row[1].text, row[2].text = r
        _fix_table_width(doc, tbl, W_COMPAT)

    _maybe_add_compat(doc, data.get("compat") or {})

    # ==== POTENCIAL / FORMATIVAS / RECOMENDACIÓN FINAL ====
    def _maybe_add_bullets(doc: Document, title: str, bullets: List[str]):
        items = [b for b in (bullets or []) if _clean_val(b)]
        if not items:
            return
        _add_section_bar(doc, title)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tbl = doc.add_table(rows=1, cols=1)
        _set_table_borders(tbl, 8, "000000")
        cell = tbl.rows[0].cells[0]
        for i, t in enumerate(items):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            _set_paragraph_compact(p)
            r0 = p.add_run("• ")
            _set_run_font(r0, bold=True)
            p.add_run(_clean_val(t))
        _set_table_fullwidth_1col(doc, tbl)

    _maybe_add_bullets(doc, "Potencial de Aporte", data.get("potencial") or [])
    _maybe_add_bullets(doc, "Recomendaciones Formativas", data.get("formativas") or [])

    rec_final = _clean_val(data.get("recomendacion_final", ""))
    if rec_final:
        _add_section_bar(doc, "Recomendación Final")
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tbl2 = doc.add_table(rows=1, cols=1)
        _set_table_borders(tbl2, 8, "000000")
        cell = tbl2.rows[0].cells[0]
        p = cell.paragraphs[0]
        _set_paragraph_compact(p)
        p.add_run(rec_final)
        _set_table_fullwidth_1col(doc, tbl2)

    # ==== FORZAR FUENTE ARIAL EN TODO EL DOCUMENTO (manteniendo centrados los títulos) ====
    FONT_NAME = "Arial"
    for p in doc.paragraphs:
        is_centered = (p.alignment == WD_ALIGN_PARAGRAPH.CENTER)
        for r in p.runs:
            r.font.name = FONT_NAME
            r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
            r._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
            r._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
            r._element.rPr.rFonts.set(qn("w:cs"), FONT_NAME)
            r.font.size = Pt(FS_BASE)
        if not is_centered:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    is_centered = (p.alignment == WD_ALIGN_PARAGRAPH.CENTER)
                    for r in p.runs:
                        r.font.name = FONT_NAME
                        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
                        r._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
                        r._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
                        r._element.rPr.rFonts.set(qn("w:cs"), FONT_NAME)
                        r.font.size = Pt(FS_BASE)
                    if not is_centered:
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ==== Guardado robusto ====
    from time import sleep
    def safe_docx_save(doc, out_path: Path, max_tries: int = 3) -> Path:
        candidate = out_path
        for i in range(max_tries):
            try:
                doc.save(str(candidate))
                return candidate
            except PermissionError:
                candidate = candidate.with_name(f"{out_path.stem}_v{i+2}{out_path.suffix}")
                sleep(0.2)
        doc.save(str(candidate))
        return candidate

    safe_docx_save(doc, out_docx)






def docx_to_pdf_optional(in_docx: Path, out_pdf: Path) -> bool:
    """
    Convierte DOCX -> PDF en Linux intentando varios binarios:
    1) soffice --headless
    2) libreoffice --headless
    3) unoconv -f pdf
    Devuelve True si el PDF se creó.
    """
    try:
        out_pdf.parent.mkdir(parents=True, exist_ok=True)
        candidates = [
            ['soffice', '--headless', '--convert-to', 'pdf', '--outdir', str(out_pdf.parent), str(in_docx)],
            ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', str(out_pdf.parent), str(in_docx)],
            ['unoconv', '-f', 'pdf', '-o', str(out_pdf), str(in_docx)],
        ]
        for cmd in candidates:
            try:
                env = os.environ.copy()
                env.setdefault('HOME', '/tmp')
                subprocess.run(cmd, env=env, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=False)
                if out_pdf.exists():
                    return True
                candidate = in_docx.with_suffix('.pdf')
                if candidate.exists():
                    try:
                        candidate.rename(out_pdf)
                    except Exception:
                        out_pdf.write_bytes(candidate.read_bytes())
                        candidate.unlink(missing_ok=True)
                    return True
            except Exception:
                pass
    except Exception:
        pass
    return False


# =========================
#   EXTRA: capturar "Tipo asistente: ..."
# =========================
TIPO_RE = re.compile(r"^\s*tipo\s*(?:de\s*)?asistente\s*:\s*(.+?)\s*$", re.IGNORECASE | re.MULTILINE)

def extract_user_tipo_asistente(txt: str) -> str:
    """Devuelve el valor literal del usuario si existe, sin limitar ni normalizar."""
    if not txt:
        return ""
    m = TIPO_RE.search(txt)
    if not m:
        return ""
    raw = m.group(1).strip()
    if (raw.startswith('"') and raw.endswith('"')) or (raw.startswith("'") and raw.endswith("'")):
        raw = raw[1:-1].strip()
    return raw


# =========================
#   MAIN
# =========================
def main():
    ap = argparse.ArgumentParser(description="Generador de Ficha desde TXT + CV con IA (DOCX estilizado)")
    ap.add_argument("--cv_text", type=str, default="")
    ap.add_argument("--cv_text_file", type=str, default="")
    ap.add_argument("--cv", type=str, default="")
    ap.add_argument("--extra", type=str, default="")
    ap.add_argument("--extra_file", type=str, default="")
    ap.add_argument("--outdir", type=str, required=True)
    ap.add_argument("--basename", type=str, required=True)
    ap.add_argument("--logo", type=str, default="")
    ap.add_argument("--no-pdf", action="store_true")
    args = ap.parse_args()

    # TXT (formulario)
    txt = args.extra.strip()
    if (not txt) and args.extra_file:
        txt = read_text(Path(args.extra_file))
    if not txt:
        raise RuntimeError("Debes pasar --extra o --extra_file con el texto de entrevista+formulario.")

    # CV
    cv_text = ""
    if args.cv_text:
        cv_text = args.cv_text
    elif args.cv_text_file:
        cv_text = read_text(Path(args.cv_text_file))
    elif args.cv:
        cv_text = read_cv_any(Path(args.cv))

    # Herramientas SOLO desde regex en TXT
    tools_regex = extract_tools_with_percent_from_text(txt, maxn=PLACE["herrs"])

    # Llamado único a IA
    data = ai_extract_all_fields(txt, cv_text, tools_regex)

    # 🔸 OVERRIDE DEL TÍTULO SEGÚN USUARIO (libre, sin confirmar)
    tipo_usuario = extract_user_tipo_asistente(txt)  # <— literal del usuario
    if tipo_usuario:
        # Guardamos campo explícito para el render
        data["titulo_usuario"] = tipo_usuario
        # También lo dejamos visible en mapping para depuración del backend
        data.setdefault("_meta", {})["titulo_origen"] = "usuario"
        data["_meta"]["titulo_usuario"] = tipo_usuario
    else:
        data.setdefault("_meta", {})["titulo_origen"] = "ia"

    # Filtros de sanidad (no agrega nada)
    _enforce_tech_areas_rules(data)

    # Validación PERMISIVA: no aborta; escribe advertencias en stderr
    warnings = _soft_validate_ai_payload(data, tools_regex)
    if warnings:
        import sys
        sys.stderr.write("[WARN] Validacion suave:\n" + "\n".join(f" - {w}" for w in warnings) + "\n")

    # Render
    outdir = Path(args.outdir); outdir.mkdir(parents=True, exist_ok=True)
    out_docx = outdir / f"{args.basename}.docx"
    out_pdf  = outdir / f"{args.basename}.pdf"

    render_docx(out_docx, data, logo_path=args.logo or None)

    used_pdf = False
    if not args.no_pdf:
        if docx2pdf_convert is not None:
            try:
                outdocx_parent = out_docx.parent
                docx2pdf_convert(str(out_docx), str(outdocx_parent))
                if out_pdf.exists():
                    used_pdf = True
                else:
                    cand = out_docx.with_suffix(".pdf")
                    if cand.exists():
                        cand.rename(out_pdf)
                        used_pdf = True
            except Exception:
                used_pdf = docx_to_pdf_optional(out_docx, out_pdf)
        else:
            used_pdf = docx_to_pdf_optional(out_docx, out_pdf)

    print(json.dumps({
        "ok": True,
        "docx": str(out_docx),
        "pdf": str(out_pdf) if used_pdf else "",
        "used_pdf": used_pdf,
        "mapping": data,  # incluye _meta.titulo_origen y, si aplica, titulo_usuario
    }, ensure_ascii=False))


if __name__ == "__main__":
    main()






